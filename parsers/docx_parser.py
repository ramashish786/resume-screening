"""
parsers/docx_parser.py
───────────────────────
Extracts text from .docx resume files using python-docx.
Handles paragraphs, tables, headers, and footers.
"""

from __future__ import annotations

import hashlib
import io

from loguru import logger

try:
    import docx
except ImportError:
    docx = None  # type: ignore


def parse_docx(file_bytes: bytes, file_name: str) -> dict:
    """
    Extract text from a .docx file.

    Args:
        file_bytes: Raw bytes of the uploaded DOCX.
        file_name:  Original filename.

    Returns:
        dict with keys: raw_text, page_count, word_count, warnings, file_hash
    """
    if docx is None:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    warnings: list[str] = []
    file_hash = hashlib.md5(file_bytes).hexdigest()

    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Could not open DOCX '{file_name}': {e}") from e

    text_parts: list[str] = []

    # Extract paragraphs (preserves heading hierarchy)
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            # Add blank line after headings for readability
            if para.style.name.startswith("Heading"):
                text_parts.append(f"\n{text}\n")
            else:
                text_parts.append(text)

    # Extract tables — preserve cell structure for skills/experience tables
    for table in document.tables:
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
        if table_rows:
            text_parts.append("\n".join(table_rows))

    # Extract header and footer text (often contains contact info)
    for section in document.sections:
        for part in (section.header, section.footer):
            try:
                hf_text = "\n".join(
                    p.text.strip() for p in part.paragraphs if p.text.strip()
                )
                if hf_text:
                    text_parts.insert(0, hf_text)  # contact info goes first
            except Exception:
                pass

    raw_text = "\n".join(text_parts)
    word_count = len(raw_text.split())

    # python-docx doesn't expose page count natively; estimate from word count
    estimated_pages = max(1, word_count // 300)

    if word_count < 30:
        warnings.append("Very little text found in DOCX. Document may be corrupt or heavily image-based.")

    logger.info(
        f"DOCX parsed: {file_name} | ~pages={estimated_pages} | words={word_count} | warnings={len(warnings)}"
    )

    return {
        "raw_text": raw_text,
        "page_count": estimated_pages,
        "word_count": word_count,
        "warnings": warnings,
        "file_hash": file_hash,
    }